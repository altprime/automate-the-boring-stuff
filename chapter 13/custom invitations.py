'''
custom invitations

text file of guest names -> guests.txt
generate word document with custom invitations as given

'''

import docx

# open guests.txt, pick names of guests.
guests = open ('guests.txt','r')
nameList = guests.read().split('\n')

docInv = docx.Document()
i = 1
for guestname in nameList:
    docInv.add_paragraph('It would be a pleasure to have the company of ',style=None)
    docInvi.add_paragraph(guestname,style=None)
    docInvi.add_paragraph('at 11010 memory lane on the evening of ',style=None)
    docInvi.add_paragraph('April 1st',style=None)
    docInvi.add_paragraph('at 7 o\'clock',style=None)
    docInvi.paragraphs[5*i-1].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)
    i += 1
    if i > 5:
        break

docInv.save('invitatoins.docx')
guests.close()
